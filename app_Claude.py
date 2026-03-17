import os
import re
import warnings
import streamlit as st
from dotenv import load_dotenv
from mistralai import Mistral
from supabase import create_client
import openpyxl

warnings.filterwarnings("ignore", category=UserWarning, module="urllib3")

# Streamlit Cloud: read from st.secrets — locally: read from .env file
try:
    MISTRAL_API_KEY      = st.secrets["MISTRAL_API_KEY"]
    SUPABASE_URL         = st.secrets["SUPABASE_URL"]
    SUPABASE_SERVICE_KEY = st.secrets["SUPABASE_SERVICE_KEY"]
except Exception:
    load_dotenv(os.path.join(os.path.dirname(__file__), "config_Claude.env"))
    MISTRAL_API_KEY      = os.getenv("MISTRAL_API_KEY")
    SUPABASE_URL         = os.getenv("SUPABASE_URL")
    SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY")

mistral  = Mistral(api_key=MISTRAL_API_KEY)
supabase = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)

YELLOW     = "FF92D050"
EXCEL_PATH = os.path.join(os.path.dirname(__file__), "themen_Claude.xlsx")

CHUNK_CHARS   = 1500    # target chunk size in characters (~500 words)
OVERLAP_CHARS = 200     # overlap between adjacent chunks

# Mistral API pricing (USD, Scale Plan — March 2026)
PRICE_OCR_PER_PAGE    = 0.002              # $2 per 1 000 pages
PRICE_EMBED_PER_TOKEN = 0.10 / 1_000_000  # $0.10 per 1M tokens
PRICE_LARGE_IN_TOKEN  = 2.0  / 1_000_000  # $2 per 1M input tokens
PRICE_LARGE_OUT_TOKEN = 6.0  / 1_000_000  # $6 per 1M output tokens


def chunk_text(text: str) -> list[str]:
    """Split text into overlapping chunks, breaking at word boundaries."""
    if len(text) <= CHUNK_CHARS:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = start + CHUNK_CHARS
        if end < len(text):
            break_at = text.rfind(' ', start + CHUNK_CHARS // 2, end)
            if break_at == -1:
                break_at = end
            end = break_at
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - OVERLAP_CHARS
        if start >= len(text):
            break
    return chunks if chunks else [text]

# Maps Excel sheet names → subject key stored in DB
SHEET_TO_SUBJECT = {
    "Themenliste Deutsch SEK II": "Deutsch",
    "Themenliste Mathe SEK II":   "Mathematik",
}


def load_indexed_books():
    """Return {subject: [filename, ...]} for all books in the DB."""
    rows = supabase.table("documents").select("filename, subject").execute()
    books = {}
    for r in rows.data:
        subj = r["subject"] or "Sonstige"
        fname = r["filename"]
        if fname not in books.get(subj, []):
            books.setdefault(subj, [])
            if fname not in books[subj]:
                books[subj].append(fname)
    return books


def load_topics_from_db():
    """Return list of (subject, course_type, topic, pinned) — pinned first."""
    rows = supabase.table("topics").select("subject, course_type, topic, pinned") \
        .order("pinned", desc=True).order("subject").order("course_type").order("topic").execute()
    return [(r["subject"], r["course_type"] or "", r["topic"], r["pinned"]) for r in rows.data]


# ── Subject normalization ───────────────────────────────────────────────────────
_SUBJECT_ALIASES = {
    "mathematik": "Mathematik",
    "mathe":      "Mathematik",
    "deutsch":    "Deutsch",
    "german":     "Deutsch",
}

def normalize_subject(raw: str) -> str:
    """Map Mistral's free-text subject string to the canonical DB key."""
    lower = raw.lower()
    for alias, canonical in _SUBJECT_ALIASES.items():
        if alias in lower:
            return canonical
    return raw.strip()


# ── Topic keyword matching ──────────────────────────────────────────────────────
_STOP = {
    "und", "oder", "der", "die", "das", "von", "in", "an", "auf", "zu", "mit",
    "für", "aus", "bei", "nach", "über", "unter", "vor", "durch", "eine", "eines",
    "einem", "einen", "sowie", "auch", "als", "ihrer", "ihren", "ihrem",
}
TOPIC_PLACEHOLDER = "inhalticher schwerpunkt / konkretes unterrichtsthema"


def _kw(text: str) -> set[str]:
    return {w for w in re.split(r'[\s,;:()\[\]→±/]+', text.lower())
            if len(w) > 3 and w not in _STOP}


def get_excel_topics_keywords(subject: str = None) -> dict[str, set[str]]:
    """Return {topic: keywords} for Excel topics (placeholder row excluded)."""
    q = supabase.table("topics").select("topic").eq("source", "excel")
    if subject:
        q = q.eq("subject", subject)
    return {
        r["topic"]: _kw(r["topic"])
        for r in q.execute().data
        if r["topic"].lower().strip() != TOPIC_PLACEHOLDER
    }


def find_matching_excel_topic(extracted: str, excel_kw: dict[str, set[str]]) -> str | None:
    """Return best-matching Excel topic when keyword overlap ≥ 55%, else None."""
    ext_kw = _kw(extracted)
    if not ext_kw:
        return None
    best_match, best_score = None, 0.0
    for excel_topic, exc_kw in excel_kw.items():
        if not exc_kw:
            continue
        overlap = len(ext_kw & exc_kw)
        min_len  = min(len(ext_kw), len(exc_kw))
        score    = overlap / min_len if min_len else 0
        if score > best_score:
            best_score, best_match = score, excel_topic
    return best_match if best_score >= 0.55 else None


DEFAULT_EXTRACTION_PROMPT = """\
Du bist Experte für NRW-Lehrpläne (Sekundarstufe II).
Analysiere den Lehrplan und gib deine Antwort exakt in diesem Format zurück:

FACH: [Fachname, z.B. Deutsch]

=== EF ===
- Thema 1
- Thema 2

=== GK ===
- Thema 1

=== LK ===
- Thema 1

Regeln:
- EF = Einführungsphase, GK = Qualifikationsphase Grundkurs, LK = Qualifikationsphase Leistungskurs
- Nur konkrete inhaltliche Schwerpunkte, keine allgemeinen Kompetenzformulierungen
- Wenn ein Kurstyp im Dokument nicht vorkommt, lass ihn weg
- Keine Nummerierung, keine Erklärungen außer den Themen selbst\
"""


def load_extraction_prompt() -> str:
    row = supabase.table("settings").select("value").eq("key", "extraction_prompt").execute()
    return row.data[0]["value"] if row.data else DEFAULT_EXTRACTION_PROMPT


def save_extraction_prompt(prompt: str):
    supabase.table("settings").upsert(
        {"key": "extraction_prompt", "value": prompt},
        on_conflict="key",
    ).execute()


def extract_topics_with_mistral(pdf_bytes: bytes, filename: str) -> tuple[str, dict]:
    """OCR a Lehrplan PDF, auto-detect subject + course types, extract topics.
    Returns (subject, {"EF": [...], "GK": [...], "LK": [...]})."""
    mfile = mistral.files.upload(
        file={"file_name": filename, "content": pdf_bytes}, purpose="ocr"
    )
    signed = mistral.files.get_signed_url(file_id=mfile.id)
    ocr = mistral.ocr.process(
        model="mistral-ocr-latest",
        document={"type": "document_url", "document_url": signed.url},
    )
    mistral.files.delete(file_id=mfile.id)

    full_text = "\n\n".join(p.markdown for p in ocr.pages if p.markdown)

    extraction_prompt = load_extraction_prompt()
    resp = mistral.chat.complete(
        model="mistral-large-latest",
        messages=[
            {"role": "system", "content": extraction_prompt},
            {"role": "user", "content": f"Lehrplan-Text:\n{full_text[:40000]}"},
        ],
    )
    raw = resp.choices[0].message.content.strip()

    # Parse subject
    subject = "Unbekannt"
    for line in raw.splitlines():
        if line.startswith("FACH:"):
            subject = normalize_subject(line.replace("FACH:", "").strip())
            break

    # Parse course type sections
    by_course: dict[str, list[str]] = {}
    current = None
    for line in raw.splitlines():
        stripped = line.strip()
        if stripped in ("=== EF ===", "=== GK ===", "=== LK ==="):
            current = stripped.replace("=", "").strip()
        elif current and stripped.startswith("-"):
            topic = stripped.lstrip("- ").strip()
            if topic:
                by_course.setdefault(current, []).append(topic)

    return subject, by_course


DEFAULT_SYSTEM_PROMPT = """\
Du bist ein Assistent für Lehrerinnen und Lehrer in Nordrhein-Westfalen (Sekundarstufe II).

Deine Aufgabe ist es, Lehrern zu helfen, passendes Unterrichtsmaterial für ihre Schülerinnen und Schüler zu erstellen. Viele Lehrer wissen auf Basis der Lehrplan-Themen allein nicht genau, welche konkreten Inhalte sie im Unterricht behandeln sollen — die folgenden Zusammenfassungen aus den Schulbüchern geben ihnen die nötige inhaltliche Grundlage dafür.

Fasse die bereitgestellten Schulbuchauszüge zum angegebenen Thema klar und strukturiert zusammen. Orientiere dich dabei an den Kompetenzerwartungen des NRW-Kernlehrplans. Nenne bei jedem wichtigen Punkt die Quelle (Dateiname und Seitenzahl).

**Stil und Format:**
- Schreibe in aktivierender, niederschwelliger Sprache — direkt und verständlich, kein akademischer Jargon
- Sprich die Leserin / den Leser mit „du" an
- Strukturiere die Zusammenfassung in mehrere Abschnitte (z. B. „Content 1", „Content 2" usw.), jeder Abschnitt behandelt einen klar abgegrenzten Teilaspekt des Themas
- Beginne jeden Abschnitt mit einer aussagekräftigen Überschrift mit passendem Emoji (z. B. ✅ 🗣️ ⚡️ 📌)
- Hebe wichtige Fachbegriffe mit **Fettdruck** oder `Backticks` hervor
- Falls ein Beispieldokument als Referenzstil beigefügt ist, orientiere dich eng an dessen Aufbau und Tonalität

Antworte auf Deutsch.\
"""


def load_system_prompt() -> str:
    row = supabase.table("settings").select("value").eq("key", "system_prompt").execute()
    return row.data[0]["value"] if row.data else DEFAULT_SYSTEM_PROMPT


def save_system_prompt(prompt: str):
    supabase.table("settings").upsert(
        {"key": "system_prompt", "value": prompt},
        on_conflict="key",
    ).execute()


# ── Example documents (few-shot style references) ──────────────────────────────

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a .docx file (no external library needed)."""
    import zipfile
    import xml.etree.ElementTree as ET
    with zipfile.ZipFile(file_bytes) as z:
        xml_content = z.read("word/document.xml")
    root = ET.fromstring(xml_content)
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    paragraphs = []
    for para in root.iter(f"{ns}p"):
        parts = [t.text for t in para.iter(f"{ns}t") if t.text]
        line = "".join(parts).strip()
        if line:
            paragraphs.append(line)
    return "\n".join(paragraphs)


def list_examples() -> list[dict]:
    """Return all uploaded example documents (without embeddings)."""
    rows = supabase.table("examples") \
        .select("id, filename, topic_name, subject, uploaded_at") \
        .order("uploaded_at", desc=True).execute()
    return rows.data or []


def delete_example(example_id: int):
    supabase.table("examples").delete().eq("id", example_id).execute()


def is_example_uploaded(filename: str) -> bool:
    row = supabase.table("examples").select("id").eq("filename", filename).limit(1).execute()
    return len(row.data) > 0


def find_closest_example(query_embedding: list, subject: str = None) -> dict | None:
    """Return the example with highest cosine similarity, or None if table is empty."""
    try:
        params = {"query_embedding": query_embedding, "match_count": 1}
        if subject:
            params["subject_filter"] = subject
        rows = supabase.rpc("match_examples", params).execute()
        if rows.data:
            return rows.data[0]
    except Exception:
        pass
    return None


def get_cached_summary(topic: str):
    """Return cached summary + sources if this topic was already run."""
    row = supabase.table("summary_cache").select("*").eq("topic", topic).execute()
    if row.data:
        # increment hit counter
        supabase.table("summary_cache").update(
            {"hits": row.data[0]["hits"] + 1}
        ).eq("topic", topic).execute()
        return row.data[0]["summary"], row.data[0]["sources"]
    return None, None


def save_cached_summary(topic: str, summary: str, sources: list):
    supabase.table("summary_cache").upsert(
        {"topic": topic, "summary": summary, "sources": sources, "hits": 1},
        on_conflict="topic",
    ).execute()


def generate_docx(topic: str, summary_text: str, sources: list) -> bytes:
    """Generate a DOCX file from a summary and return it as bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    import re
    import io

    doc = Document()

    # Title
    title = doc.add_heading(topic, level=0)
    title.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    doc.add_paragraph()

    # Parse and add summary — handle basic Markdown
    for line in summary_text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            p = doc.add_paragraph()
            # Handle **bold** inline
            parts = re.split(r"\*\*(.+?)\*\*", stripped)
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True

    # Sources section
    if sources:
        doc.add_paragraph()
        doc.add_heading("Quellen", level=2)
        for r in sources:
            chunk_label = f" (Abschnitt {r['chunk_index']+1})" if r.get("chunk_index", 0) > 0 else ""
            doc.add_paragraph(
                f"{r['filename']} — Seite {r['page_number']}{chunk_label} "
                f"(Relevanz: {r['similarity']:.0%})",
                style="List Bullet",
            )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def is_already_indexed(filename: str) -> bool:
    row = supabase.table("documents").select("id").eq("filename", filename).limit(1).execute()
    return len(row.data) > 0


def log_processing_cost(filename: str, operation: str, pages: int = None,
                        tokens_in: int = None, tokens_out: int = None, cost_usd: float = 0.0):
    """Append one row to processing_log. Never raises — cost tracking must not block the pipeline."""
    try:
        supabase.table("processing_log").insert({
            "filename": filename, "operation": operation,
            "pages": pages, "tokens_in": tokens_in,
            "tokens_out": tokens_out, "cost_usd": round(cost_usd, 6),
        }).execute()
    except Exception:
        pass


def get_book_costs(filename: str) -> dict:
    """Return {total_cost_usd, pages} aggregated from processing_log for a book."""
    try:
        rows = supabase.table("processing_log") \
            .select("operation, pages, cost_usd") \
            .eq("filename", filename).execute()
        if not rows.data:
            return {}
        total = sum(r["cost_usd"] or 0 for r in rows.data)
        pages = sum(r["pages"] or 0 for r in rows.data if r.get("operation") == "ocr")
        return {"total_cost_usd": total, "pages": pages}
    except Exception:
        return {}


def load_lehrplan_from_cache(filename: str):
    """Return (subject, by_course_dict) if this PDF was already extracted, else (None, None)."""
    rows = supabase.table("topics").select("subject, course_type, topic") \
        .eq("source", "lehrplan").eq("source_file", filename).execute().data
    if not rows:
        return None, None
    subject = rows[0]["subject"]
    by_course: dict[str, list[str]] = {}
    for r in rows:
        by_course.setdefault(r["course_type"] or "EF", []).append(r["topic"])
    return subject, by_course




# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(page_title="student PRO — PDF Engine", layout="wide")
st.title("student PRO — PDF Knowledge Engine")

tab1, tab_lehrplan, tab_beispiele, tab2, tab4 = st.tabs([
    "📚 Bücher hochladen", "📄 Lehrplan hochladen", "📝 Beispiele hochladen",
    "🔍 Thema abfragen", "❓ Wie funktioniert es?"
])

# ── Tab 1: Upload ──────────────────────────────────────────────────────────────
with tab1:
    st.header("PDF hochladen & indexieren")
    uploaded_file = st.file_uploader("PDF-Datei auswählen", type="pdf")

    if uploaded_file:
        st.write(f"**Datei:** {uploaded_file.name}  ({uploaded_file.size / 1_000_000:.1f} MB)")

        # ── Cache check for ingestion ──────────────────────────────────────────
        already_indexed = is_already_indexed(uploaded_file.name)
        if already_indexed:
            st.success(
                f"✅ **'{uploaded_file.name}'** ist bereits indexiert — "
                f"kein erneuter OCR-Lauf nötig. Mistral-Credits gespart."
            )
            if st.button("🔄 Trotzdem neu verarbeiten (überschreiben)"):
                already_indexed = False  # allow re-run if explicitly requested

        if not already_indexed and st.button("Buch verarbeiten", type="primary"):
            file_bytes = uploaded_file.read()
            try:
                with st.spinner("Hochladen zu Mistral OCR..."):
                    mistral_file = mistral.files.upload(
                        file={"file_name": uploaded_file.name, "content": file_bytes},
                        purpose="ocr",
                    )
                    signed = mistral.files.get_signed_url(file_id=mistral_file.id)

                with st.spinner("OCR läuft (kann 1–2 Minuten dauern)..."):
                    ocr_response = mistral.ocr.process(
                        model="mistral-ocr-latest",
                        document={"type": "document_url", "document_url": signed.url},
                    )
                    pages = ocr_response.pages

                ocr_cost = len(pages) * PRICE_OCR_PER_PAGE
                log_processing_cost(uploaded_file.name, "ocr", pages=len(pages), cost_usd=ocr_cost)
                st.info(f"OCR abgeschlossen — {len(pages)} Seiten erkannt  (${ocr_cost:.4f})")

                progress = st.progress(0, text="Starte Einbettung...")
                skipped = 0
                total_chunks_stored = 0
                total_embed_tokens = 0
                for i, page in enumerate(pages):
                    text = page.markdown.strip() if page.markdown else ""
                    if not text:
                        skipped += 1
                        progress.progress((i + 1) / len(pages), text=f"Seite {i+1}/{len(pages)} — leer, übersprungen")
                        continue

                    page_chunks = chunk_text(text[:8000])
                    texts_to_embed = [c for c in page_chunks]
                    emb_resp = mistral.embeddings.create(model="mistral-embed", inputs=texts_to_embed)
                    if hasattr(emb_resp, "usage") and emb_resp.usage:
                        total_embed_tokens += getattr(emb_resp.usage, "prompt_tokens", 0) or 0
                    for ci, chunk_content in enumerate(page_chunks):
                        embedding = emb_resp.data[ci].embedding
                        supabase.table("documents").upsert(
                            {"filename": uploaded_file.name, "page_number": page.index + 1,
                             "chunk_index": ci, "content": chunk_content, "embedding": embedding},
                            on_conflict="filename,page_number,chunk_index",
                        ).execute()
                    total_chunks_stored += len(page_chunks)

                    progress.progress((i + 1) / len(pages), text=f"Seite {i+1}/{len(pages)} ({len(page_chunks)} Abschnitte) gespeichert")

                embed_cost = total_embed_tokens * PRICE_EMBED_PER_TOKEN
                if total_embed_tokens > 0:
                    log_processing_cost(uploaded_file.name, "embed",
                                        tokens_in=total_embed_tokens, cost_usd=embed_cost)
                mistral.files.delete(file_id=mistral_file.id)
                total_cost = ocr_cost + embed_cost
                st.success(
                    f"Fertig! {len(pages) - skipped} Seiten → {total_chunks_stored} Abschnitte indexiert "
                    f"({skipped} leere Seiten übersprungen) — '{uploaded_file.name}'\n\n"
                    f"💰 **Kosten dieser Verarbeitung: ${total_cost:.4f}** "
                    f"(OCR: ${ocr_cost:.4f} · Embedding: ${embed_cost:.4f})"
                )

            except Exception as e:
                st.error(f"Fehler: {e}")
                raise

    st.divider()
    st.subheader("Indexierte Bücher")
    try:
        rows = supabase.table("documents").select("filename, page_number").execute()
        if rows.data:
            pages_per_file: dict[str, set] = {}
            chunks_per_file: dict[str, int] = {}
            for r in rows.data:
                pages_per_file.setdefault(r["filename"], set()).add(r["page_number"])
                chunks_per_file[r["filename"]] = chunks_per_file.get(r["filename"], 0) + 1
            for fname in sorted(pages_per_file.keys()):
                n_pages  = len(pages_per_file[fname])
                n_chunks = chunks_per_file[fname]
                costs = get_book_costs(fname)
                cost_label = (
                    f"  ·  💰 ${costs['total_cost_usd']:.4f}"
                    if costs.get("total_cost_usd") else ""
                )
                if n_chunks > n_pages:
                    st.write(f"- **{fname}** — {n_pages} Seiten ({n_chunks} Abschnitte){cost_label}")
                else:
                    st.write(f"- **{fname}** — {n_pages} Seiten{cost_label}")
        else:
            st.write("Noch keine Bücher indexiert.")
    except Exception as e:
        st.warning(f"Bücherliste konnte nicht geladen werden: {e}")

# ── Tab Lehrplan: Upload & extract topics ──────────────────────────────────────
with tab_lehrplan:
    st.header("Lehrplan hochladen & Themen extrahieren")

    uploaded_lehrplan = st.file_uploader("Lehrplan-PDF auswählen", type="pdf", key="lehrplan_uploader")

    # ── System-Prompt editor ───────────────────────────────────────────────────
    with st.expander("⚙️ System-Prompt anpassen"):
        current_extraction_prompt = load_extraction_prompt()
        edited_extraction_prompt = st.text_area(
            "System-Prompt für Themen-Extraktion (wird bei jedem Lehrplan-Upload an Mistral übergeben):",
            value=current_extraction_prompt,
            height=250,
            key="extraction_prompt_editor",
        )
        if st.button("💾 Prompt speichern", key="save_extraction_prompt"):
            save_extraction_prompt(edited_extraction_prompt)
            st.success("Gespeichert.")

    if uploaded_lehrplan:
        # ── Cache check ────────────────────────────────────────────────────────
        cached_subject, cached_by_course = load_lehrplan_from_cache(uploaded_lehrplan.name)
        if cached_by_course:
            st.info("💾 **Aus Cache geladen** — keine Mistral-Credits verbraucht.")
            st.session_state["extracted_by_course"] = cached_by_course
            st.session_state["extracted_subject"]   = cached_subject
            st.session_state["extracted_filename"]  = uploaded_lehrplan.name
            total = sum(len(v) for v in cached_by_course.values())
            course_summary = ", ".join(f"{ct}: {len(ts)}" for ct, ts in cached_by_course.items())
            st.success(f"Fach: **{cached_subject}** — {total} Themen ({course_summary})")
            if st.button("🔄 Neu extrahieren (Cache überschreiben)"):
                # Delete existing lehrplan rows for this file so re-extraction saves fresh
                supabase.table("topics").delete() \
                    .eq("source", "lehrplan").eq("source_file", uploaded_lehrplan.name).execute()
                st.session_state["extracted_by_course"] = {}
                st.rerun()
        elif st.button("Themen extrahieren", type="primary"):
            try:
                with st.spinner("OCR läuft..."):
                    pdf_bytes = uploaded_lehrplan.read()
                with st.spinner("Mistral erkennt Fach, Kurstyp und extrahiert Themen..."):
                    detected_subject, by_course = extract_topics_with_mistral(pdf_bytes, uploaded_lehrplan.name)
                total = sum(len(v) for v in by_course.values())
                st.session_state["extracted_by_course"] = by_course
                st.session_state["extracted_subject"]   = detected_subject
                st.session_state["extracted_filename"]  = uploaded_lehrplan.name
                course_summary = ", ".join(f"{ct}: {len(ts)}" for ct, ts in by_course.items())
                st.success(f"Fach erkannt: **{detected_subject}** — {total} Themen ({course_summary})")
            except Exception as e:
                st.error(f"Fehler: {e}")
                raise

    if st.session_state.get("extracted_by_course"):
        by_course_now = st.session_state["extracted_by_course"]

        # ── Quality metrics (keyword-overlap matching) ─────────────────────────
        detected_subj = st.session_state.get("extracted_subject")
        excel_kw = get_excel_topics_keywords(subject=detected_subj)
        all_extracted = [
            t for tlist in by_course_now.values() for t in tlist
            if t.lower().strip() != TOPIC_PLACEHOLDER
        ]
        # For each extracted topic find best-matching Excel topic
        matched_pairs: dict[str, str] = {}   # excel_topic → extracted_topic
        for ext in all_extracted:
            match = find_matching_excel_topic(ext, excel_kw)
            if match and match not in matched_pairs:
                matched_pairs[match] = ext
        matched_extracted = set(matched_pairs.values())

        col_a, col_b, col_c = st.columns(3)
        col_a.metric("Extrahierte Themen", len(all_extracted))
        col_b.metric("Themen aus Excel", len(excel_kw))
        col_c.metric("Übereinstimmungen", f"{len(matched_pairs)} / {len(excel_kw)}")
        if matched_pairs:
            with st.expander(f"✓ {len(matched_pairs)} Übereinstimmungen anzeigen"):
                for excel_t, ext_t in sorted(matched_pairs.items()):
                    st.markdown(f"**Excel:** {excel_t}  \n→ *Lehrplan:* {ext_t}")
                    st.divider()

        st.subheader("Themen auswählen")
        st.caption("Alle Themen, die du bestätigst, werden in den Dropdown übernommen. "
                   "✓ = auch in der Excel-Liste.")
        selected_new = []  # list of (course_type, topic)
        for ct, topics_list in by_course_now.items():
            st.markdown(f"**{ct}** — {len(topics_list)} Themen")
            for t in topics_list:
                if t.lower().strip() == TOPIC_PLACEHOLDER:
                    continue
                label = f"✓ {t}" if t in matched_extracted else t
                if st.checkbox(label, value=True, key=f"new_topic_{ct}_{t}"):
                    selected_new.append((ct, t))

        if st.button("✅ Ausgewählte Themen speichern", disabled=not selected_new):
            subj  = st.session_state["extracted_subject"]
            fname = st.session_state.get("extracted_filename", "")
            pinned_keys = {
                (r["topic"], r["subject"], r["course_type"])
                for r in supabase.table("topics").select("topic, subject, course_type")
                    .eq("pinned", True).execute().data
            }
            for ct, t in selected_new:
                if (t, subj, ct) in pinned_keys:
                    supabase.table("topics").update({"source_file": fname}) \
                        .eq("topic", t).eq("subject", subj).eq("course_type", ct).execute()
                else:
                    supabase.table("topics").upsert(
                        {"topic": t, "subject": subj, "course_type": ct,
                         "pinned": False, "source": "lehrplan", "source_file": fname},
                        on_conflict="topic,subject,course_type",
                    ).execute()
            # Mark matched Excel topics as in_lehrplan=True
            for excel_topic in matched_pairs.keys():
                supabase.table("topics").update({"in_lehrplan": True}) \
                    .eq("topic", excel_topic).eq("source", "excel").execute()
            st.success(f"{len(selected_new)} Themen gespeichert — {len(matched_pairs)} Excel-Themen als 'im Lehrplan' markiert.")
            st.session_state["extracted_by_course"] = {}

    st.divider()

    # ── Verwendete Lehrplan-PDFs ───────────────────────────────────────────────
    st.subheader("Verwendete Lehrplan-PDFs")
    lehrplan_files = supabase.table("topics").select("source_file, subject") \
        .eq("source", "lehrplan").execute().data
    if lehrplan_files:
        file_counts = {}
        for r in lehrplan_files:
            key = (r["source_file"] or "Unbekannt", r["subject"])
            file_counts[key] = file_counts.get(key, 0) + 1
        for (fname, subj), count in sorted(file_counts.items()):
            st.write(f"- **{fname}** ({subj}) — {count} Themen extrahiert")
    else:
        st.write("Noch keine Lehrplan-PDFs verarbeitet.")

    st.divider()

    # ── Gespeicherte Themen ────────────────────────────────────────────────────
    st.subheader("Gespeicherte Themen")
    all_topics = supabase.table("topics").select("subject, course_type, topic, pinned, in_lehrplan") \
        .order("pinned", desc=True).order("subject").order("course_type").order("topic").execute().data
    if all_topics:
        # Group by subject → course_type
        by_subject: dict[str, dict[str, list]] = {}
        for r in all_topics:
            subj = r["subject"]
            ct   = r["course_type"] or "Sonstige"
            by_subject.setdefault(subj, {}).setdefault(ct, []).append(r)

        CT_ORDER = ["EF", "GK", "LK", "Sonstige"]
        for subj in sorted(by_subject.keys()):
            ct_dict = by_subject[subj]
            total   = sum(len(v) for v in ct_dict.values())
            n_pin   = sum(1 for v in ct_dict.values() for r in v if r["pinned"])
            n_lp    = sum(1 for v in ct_dict.values() for r in v if r.get("in_lehrplan") and not r["pinned"])
            with st.expander(f"{subj} — {total} Themen ({n_pin} markiert ★, {n_lp} im Kernlehrplan)"):
                for ct in CT_ORDER:
                    if ct not in ct_dict:
                        continue
                    st.markdown(f"**{ct}**")
                    for r in ct_dict[ct]:
                        if r["pinned"]:
                            st.markdown(
                                f'<div style="background-color:#ff4b4b; color:white; '
                                f'padding:4px 12px; border-radius:4px; margin:2px 0; '
                                f'font-size:0.9em;">★ {r["topic"]}</div>',
                                unsafe_allow_html=True,
                            )
                        elif r.get("in_lehrplan"):
                            st.markdown(
                                f'<div style="background-color:#21a354; color:white; '
                                f'padding:4px 12px; border-radius:4px; margin:2px 0; '
                                f'font-size:0.9em;">✓ {r["topic"]}</div>',
                                unsafe_allow_html=True,
                            )
                        else:
                            st.write(f"　{r['topic']}")
    else:
        st.write("Noch keine Themen gespeichert.")


# ── Tab Beispiele: Upload style-reference documents ────────────────────────────
with tab_beispiele:
    st.header("Beispieldokumente hochladen")
    st.caption(
        "Lade hier Beispieldokumente hoch (docx oder pdf). "
        "Das System findet beim Abfragen automatisch das passendste Beispiel "
        "und gibt es Mistral als Stilvorlage mit — so orientiert sich die Zusammenfassung "
        "am gewünschten Format (Content 1 / Content 2 usw.)."
    )

    with st.expander("❓ Was passiert, wenn ich eine Datei hochlade?"):
        st.markdown("""
**Schritt 1 — Text lesen**
Bei einer `.docx`-Datei liest das System den Text direkt aus der Datei.
Bei einer `.pdf`-Datei schickt es die Datei an Mistral OCR, das die Seiten wie ein Scanner in Text umwandelt.

**Schritt 2 — Steckbrief berechnen**
Mistral Embed berechnet einen „Steckbrief" des Textes — eine Liste von 1024 Zahlen,
die beschreiben, worum es in dem Dokument geht. Ähnliche Themen → ähnliche Zahlen.

**Schritt 3 — Speichern**
Text + Steckbrief werden in der Datenbank abgelegt (Tabelle `examples`).

---

**Was passiert beim Abfragen?**

Wenn du ein Thema abfragst (z.B. *„Lyrische Texte"*), berechnet das System auch für dieses Thema einen Steckbrief.
Es vergleicht ihn dann mit den Steckbriefen aller hochgeladenen Beispieldokumente.

Wenn ein Beispiel gut passt (Ähnlichkeit ≥ 50 %), wird es Mistral mitgeschickt:
*„Hier ist ein Beispiel, wie das Ergebnis aussehen soll — orientiere dich an Aufbau und Stil."*

Mistral schreibt dann die Zusammenfassung im gewünschten Format (Content 1, Content 2 usw.).

---

**❓ Ändert sich der System-Prompt, wenn ich Beispiele hochlade?**

**Nein.** Der System-Prompt (unter ⚙️ im Tab „Thema abfragen") bleibt unverändert.
Das Beispieldokument wird *zusätzlich* an Mistral geschickt — als Anhang zusammen mit den Schulbuch-Auszügen.
Mistral sieht also auf einmal: Thema + Buchseiten + Stilvorlage.

Du kannst den System-Prompt trotzdem jederzeit manuell anpassen — Änderungen dort gelten dauerhaft.
        """)

    uploaded_examples = st.file_uploader(
        "Beispieldokumente auswählen (.docx oder .pdf) — mehrere gleichzeitig möglich",
        type=["docx", "pdf"],
        accept_multiple_files=True,
        key="example_uploader",
    )

    if uploaded_examples:
        new_files = [f for f in uploaded_examples if not is_example_uploaded(f.name)]
        already_uploaded = [f for f in uploaded_examples if is_example_uploaded(f.name)]

        if already_uploaded:
            st.info(f"ℹ️ {len(already_uploaded)} Datei(en) bereits vorhanden: {', '.join(f.name for f in already_uploaded)}")

        if new_files:
            st.write(f"**{len(new_files)} neue Datei(en) bereit zum Verarbeiten:**")
            for f in new_files:
                st.write(f"　· {f.name}")

        if st.button(
            f"Alle verarbeiten & speichern ({len(new_files)} neu)" if new_files else "Alle erneut hochladen (überschreiben)",
            type="primary",
            key="process_examples",
            disabled=not uploaded_examples,
        ):
            import io
            files_to_process = new_files if new_files else uploaded_examples
            progress = st.progress(0, text="Starte...")
            for idx, uploaded_file in enumerate(files_to_process):
                progress.progress(idx / len(files_to_process), text=f"Verarbeite {uploaded_file.name} ({idx+1}/{len(files_to_process)})...")
                file_bytes = uploaded_file.read()
                try:
                    if uploaded_file.name.lower().endswith(".docx"):
                        content = extract_text_from_docx(io.BytesIO(file_bytes))
                    else:
                        mfile = mistral.files.upload(
                            file={"file_name": uploaded_file.name, "content": file_bytes},
                            purpose="ocr",
                        )
                        signed = mistral.files.get_signed_url(file_id=mfile.id)
                        ocr = mistral.ocr.process(
                            model="mistral-ocr-latest",
                            document={"type": "document_url", "document_url": signed.url},
                        )
                        mistral.files.delete(file_id=mfile.id)
                        content = "\n\n".join(p.markdown for p in ocr.pages if p.markdown)

                    if not content.strip():
                        st.warning(f"⚠️ {uploaded_file.name} — kein Text extrahiert, übersprungen.")
                        continue

                    emb_resp = mistral.embeddings.create(
                        model="mistral-embed",
                        inputs=[content[:8000]],
                    )
                    embedding = emb_resp.data[0].embedding

                    lower_name = uploaded_file.name.lower()
                    detected_subject = "Deutsch" if "deutsch" in lower_name else \
                                       "Mathematik" if any(x in lower_name for x in ["mathe", "math"]) else None
                    raw_topic = uploaded_file.name.replace(".docx", "").replace(".pdf", "")

                    supabase.table("examples").upsert(
                        {"filename": uploaded_file.name, "topic_name": raw_topic,
                         "subject": detected_subject, "content": content, "embedding": embedding},
                        on_conflict="filename",
                    ).execute()
                    st.success(f"✅ {uploaded_file.name} gespeichert (Fach: {detected_subject or 'unbekannt'})")

                except Exception as e:
                    st.error(f"Fehler bei {uploaded_file.name}: {e}")

            progress.progress(1.0, text="Fertig!")
            st.rerun()

    st.divider()
    st.subheader("Gespeicherte Beispieldokumente")
    examples = list_examples()
    if not examples:
        st.write("Noch keine Beispiele hochgeladen.")
    else:
        for ex in examples:
            col_a, col_b = st.columns([5, 1])
            with col_a:
                subj_label = f"  ·  Fach: {ex['subject']}" if ex["subject"] else ""
                st.write(f"**{ex['filename']}**{subj_label}")
            with col_b:
                if st.button("🗑️ Löschen", key=f"del_example_{ex['id']}"):
                    delete_example(ex["id"])
                    st.rerun()

    # ── Supabase unique constraint for examples.filename ──────────────────────
    # (Created once via Management API — documented here for reference)
    # ALTER TABLE examples ADD CONSTRAINT examples_filename_key UNIQUE (filename);


# ── Tab 2: Search ──────────────────────────────────────────────────────────────
with tab2:
    st.header("Thema abfragen")

    # ── Book selector ──────────────────────────────────────────────────────────
    st.subheader("Bücher auswählen")
    indexed_books = load_indexed_books()
    selected_books = []
    if not indexed_books:
        st.warning("Noch keine Bücher indexiert.")
    else:
        for subj in sorted(indexed_books.keys()):
            st.markdown(f"**{subj}**")
            for fname in sorted(indexed_books[subj]):
                checked = st.checkbox(fname, value=True, key=f"book_{fname}")
                if checked:
                    selected_books.append(fname)
        if not selected_books:
            st.warning("Bitte mindestens ein Buch auswählen.")
    st.divider()

    options = load_topics_from_db()  # list of (subject, course_type, topic, pinned)

    if not options:
        st.warning("Noch keine Themen geladen. Bitte zuerst einen Lehrplan hochladen.")
    else:
        col1, = st.columns([1])
        with col1:
            labels = [
                f"{'★ ' if pinned else ''}{topic}  [{subject} · {course_type}]"
                for subject, course_type, topic, pinned in options
            ]
            selected_idx = st.selectbox(
                "Thema auswählen:",
                range(len(labels)),
                format_func=lambda i: labels[i],
            )
            subject     = options[selected_idx][0]
            course_type = options[selected_idx][1]
            keyword     = options[selected_idx][2]
        top_k = 20

        # ── System-Prompt Editor ───────────────────────────────────────────────
        with st.expander("⚙️ System-Prompt anpassen"):
            current_prompt = load_system_prompt()
            edited_prompt = st.text_area(
                "System-Prompt (wird bei jeder Zusammenfassung an Mistral übergeben):",
                value=current_prompt,
                height=220,
                key="system_prompt_editor",
            )
            if st.button("💾 Prompt speichern"):
                save_system_prompt(edited_prompt)
                st.success("Gespeichert.")

        # ── Cache check for summary ────────────────────────────────────────────
        if "auto_generate" not in st.session_state:
            st.session_state["auto_generate"] = False

        cached_summary, cached_sources = get_cached_summary(keyword)
        if cached_summary:
            if st.session_state.get("fresh_topic") == keyword:
                st.success("✅ Neu generiert")
            else:
                st.info("💾 **Aus Cache geladen** — keine Mistral-Credits verbraucht.")
            with st.expander("📝 Zusammenfassung", expanded=True):
                st.markdown(cached_summary)
            if cached_sources:
                with st.expander(f"📚 Quellseiten ({len(cached_sources)} Treffer)", expanded=False):
                    for r in cached_sources:
                        chunk_label = f" (Abschnitt {r['chunk_index']+1})" if r.get('chunk_index', 0) > 0 else ""
                        with st.expander(f"**{r['filename']}** — Seite {r['page_number']}{chunk_label}  (Relevanz: {r['similarity']:.0%})"):
                            st.write(r["content"][:800] + ("..." if len(r["content"]) > 800 else ""))

            docx_bytes = generate_docx(keyword, cached_summary, cached_sources or [])
            st.download_button(
                "⬇️ Als Word-Dokument herunterladen (.docx)",
                data=docx_bytes,
                file_name=f"{keyword[:60].replace('/', '-')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

            if st.button("🔄 Neu generieren (Cache überschreiben)"):
                supabase.table("summary_cache").delete().eq("topic", keyword).execute()
                st.session_state["auto_generate"] = True
                st.rerun()

        run_generation = (
            not cached_summary and
            (st.session_state.get("auto_generate") or
             st.button("Relevante Inhalte abrufen", type="primary", disabled=not selected_books))
        )
        if run_generation:
            st.session_state["auto_generate"] = False
            try:
                with st.spinner("Thema wird analysiert (Steckbrief wird berechnet)..."):
                    emb = mistral.embeddings.create(model="mistral-embed", inputs=[keyword])
                    query_embedding = emb.data[0].embedding
                    query_embed_tokens = getattr(emb.usage, "prompt_tokens", 0) if hasattr(emb, "usage") and emb.usage else 0

                books_label = ", ".join(b.split(".")[0] for b in selected_books)
                with st.spinner(f"Bücher werden durchsucht: {books_label}..."):
                    result = supabase.rpc(
                        "match_documents",
                        {"query_embedding": query_embedding, "match_count": int(top_k),
                         "subject_filter": subject,
                         "filename_filter": selected_books},
                    ).execute()
                    chunks = result.data

                if not chunks:
                    st.warning("Keine relevanten Seiten gefunden. Bitte zuerst ein Buch hochladen.")
                else:
                    # Show which books contributed how many chunks
                    book_hits: dict[str, int] = {}
                    for r in chunks:
                        book_hits[r["filename"]] = book_hits.get(r["filename"], 0) + 1
                    hits_label = " · ".join(
                        f"{fn.split('.')[0]} ({n} Abschnitte)" for fn, n in book_hits.items()
                    )
                    st.caption(f"📖 Gefunden in: {hits_label}")

                    context = "\n\n---\n\n".join(
                        [f"[{r['filename']}, Seite {r['page_number']}]\n{r['content']}" for r in chunks]
                    )

                    # ── Find closest example for style reference ───────────────
                    closest_example = find_closest_example(query_embedding, subject=subject)
                    all_examples = list_examples()
                    example_names = ", ".join(f"**{e['filename']}**" for e in all_examples) if all_examples else "—"

                    if closest_example:
                        example_block = (
                            f"\n\n---\n\n**Beispieldokument als Stilvorlage** "
                            f"(orientiere dich an Aufbau und Tonalität, nicht am Inhalt):\n\n"
                            f"{closest_example['content'][:4000]}"
                        )
                        example_note = f"📄 Stilvorlage verwendet: **{closest_example['filename']}**"
                        st.caption(f"📄 Stilvorlagen vorhanden: {example_names} — verwendet: **{closest_example['filename']}**")
                    else:
                        example_block = ""
                        example_note = None
                        st.caption("📄 Keine Stilvorlagen hochgeladen")

                    with st.spinner("Zusammenfassung wird erstellt (Mistral Large)..."):
                        system_prompt = load_system_prompt()
                        user_message = (
                            f"Thema: {keyword}\n\nRelevante Auszüge:\n{context}"
                            f"{example_block}"
                        )
                        response = mistral.chat.complete(
                            model="mistral-large-latest",
                            messages=[
                                {"role": "system", "content": system_prompt},
                                {"role": "user", "content": user_message},
                            ],
                        )

                    # ── Cost calculation ───────────────────────────────────────
                    chat_in  = response.usage.prompt_tokens     if hasattr(response, "usage") and response.usage else 0
                    chat_out = response.usage.completion_tokens if hasattr(response, "usage") and response.usage else 0
                    query_cost = query_embed_tokens * PRICE_EMBED_PER_TOKEN
                    chat_cost  = chat_in * PRICE_LARGE_IN_TOKEN + chat_out * PRICE_LARGE_OUT_TOKEN
                    total_query_cost = query_cost + chat_cost
                    st.info(
                        f"💰 Kosten dieser Abfrage: **${total_query_cost:.4f}** "
                        f"(Steckbrief: ${query_cost:.5f} · "
                        f"Zusammenfassung: ${chat_cost:.4f} — "
                        f"{chat_in} Input-Tokens, {chat_out} Output-Tokens)"
                    )

                    summary_text = response.choices[0].message.content
                    sources = [
                        {"filename": r["filename"], "page_number": r["page_number"],
                         "similarity": r["similarity"], "content": r["content"]}
                        for r in chunks
                    ]

                    # Save to cache + mark as freshly generated
                    save_cached_summary(keyword, summary_text, sources)
                    st.session_state["fresh_topic"] = keyword

                    with st.expander("📝 Zusammenfassung", expanded=True):
                        if example_note:
                            st.caption(example_note)
                        st.markdown(summary_text)

                    docx_bytes = generate_docx(keyword, summary_text, sources)
                    st.download_button(
                        "⬇️ Als Word-Dokument herunterladen (.docx)",
                        data=docx_bytes,
                        file_name=f"{keyword[:60].replace('/', '-')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

                    with st.expander(f"📚 Quellseiten ({len(chunks)} Treffer)", expanded=False):
                        for r in chunks:
                            chunk_label = f" (Abschnitt {r.get('chunk_index', 0)+1})" if r.get('chunk_index', 0) > 0 else ""
                            with st.expander(
                                f"**{r['filename']}** — Seite {r['page_number']}{chunk_label}  "
                                f"(Relevanz: {r['similarity']:.0%})"
                            ):
                                st.write(r["content"][:800] + ("..." if len(r["content"]) > 800 else ""))

            except Exception as e:
                st.error(f"Fehler bei der Suche: {e}")
                raise



# ── Tab 4: How it works ────────────────────────────────────────────────────────
with tab4:
    st.header("❓ Wie funktioniert das System?")
    st.caption("Hier erklären wir alles — so einfach wie möglich, ohne Fachwissen nötig.")

    st.markdown("---")

    st.subheader("🧠 Die Grundidee — in einem Satz")
    st.markdown("""
Du gibst dem System Schulbücher. Das System liest sie komplett durch, merkt sich jeden Abschnitt,
und kann später auf Knopfdruck alles finden, was zu einem bestimmten Lehrplan-Thema passt —
und schreibt daraus automatisch eine Zusammenfassung.
    """)

    st.markdown("---")

    with st.expander("📚 Schritt 1 — Bücher hochladen (einmalig)"):
        st.markdown("""
**Was du tust:** Du lädst ein Schulbuch als PDF hoch.

**Was das System macht:**

1. **Lesen** — Die KI (Mistral OCR) liest jede einzelne Seite des PDFs und wandelt sie in normalen Text um.
   Das ist nötig, weil PDFs oft keine einfachen Textdateien sind — besonders bei eingescannten Büchern.

2. **Steckbrief erstellen** — Für jeden Textabschnitt (~500 Wörter) wird ein sogenannter *Steckbrief* berechnet.
   Das ist eine Liste von 1024 Zahlen, die beschreiben, worum es in diesem Abschnitt geht.
   Ähnliche Themen bekommen ähnliche Steckbriefe — das ist die Magie dahinter.

3. **Speichern** — Text + Steckbrief + Seitenzahl werden in der Datenbank abgelegt.

**Wie oft?** Nur einmal pro Buch. Danach ist alles gespeichert und das Buch muss nie wieder verarbeitet werden.

**Was kostet das?** Rund **0,20 € pro 100 Seiten** — also etwa 0,60 € für ein typisches 300-Seiten-Buch.
Das Lesen (OCR) macht fast den ganzen Betrag aus, die Steckbriefe sind winzig.
        """)

    with st.expander("📄 Schritt 2 — Lehrplan hochladen (einmalig)"):
        st.markdown("""
**Was du tust:** Du lädst den NRW-Kernlehrplan als PDF hoch.

**Was das System macht:**

1. Die KI liest den gesamten Lehrplan-Text
2. Sie erkennt automatisch das Fach (Deutsch, Mathe, ...) und ob es EF, GK oder LK ist
3. Sie extrahiert alle konkreten Themen — z.B. *„Sprachvarietäten und ihre gesellschaftliche Bedeutung"*
4. Du siehst die extrahierten Themen und kannst auswählen, welche gespeichert werden sollen

**Wozu?** Die gespeicherten Themen erscheinen später als Dropdown im Tab „Thema abfragen" —
so musst du nicht jedes Mal den genauen Thementitel eintippen.

**Markierungen in der Liste:**
- ⭐ **Rot** — persönlich markierte Themen (höchste Priorität)
- ✓ **Grün** — Themen, die sowohl in der Excel-Liste als auch im Kernlehrplan stehen
- **Ohne Markierung** — nur im Kernlehrplan
        """)

    with st.expander("🔍 Schritt 3 — Thema abfragen (jederzeit)"):
        st.markdown("""
**Was du tust:** Du wählst ein Thema aus dem Dropdown aus und klickst auf „Relevante Inhalte abrufen".

**Was das System macht — Schritt für Schritt:**

1. **Dein Thema bekommt auch einen Steckbrief** — genau wie die Buchseiten.

2. **Vergleich** — Das System vergleicht den Steckbrief deines Themas mit den Steckbriefen
   aller gespeicherten Buchseiten. Je ähnlicher die Steckbriefe, desto relevanter die Seite.

3. **Top 10** — Die 10 ähnlichsten Abschnitte werden ausgewählt. Das sind die Fundstellen.

4. **Zusammenfassung schreiben** — Die 10 Abschnitte werden zusammen an Mistral Large geschickt,
   zusammen mit der Anweisung: *„Schreib eine strukturierte Zusammenfassung für Lehrer in NRW."*
   Mistral schreibt dann den fertigen Text.

5. **Gespeichert** — Das Ergebnis wird gecacht. Wenn jemand dasselbe Thema nochmal abfragt,
   kommt die gespeicherte Version zurück — ohne nochmal Mistral zu bezahlen.

**Was sind die „Quellseiten" unten?**
Das sind die 10 Abschnitte, die das System ausgewählt hat — mit Buchtitel, Seitenzahl und
Ähnlichkeitsprozentsatz. Du kannst jeden aufklappen und den Originaltext lesen.
        """)

    with st.expander("📝 Schritt 4 — Beispieldokumente hochladen (optional, aber wichtig)"):
        st.markdown("""
**Was das ist:** Fertig geschriebene Beispieltexte — so wie die Zusammenfassungen
am Ende aussehen sollen. Mit Kapiteln (Content 1, Content 2 usw.) im gewünschten Stil.

**Was du tust:** Du lädst diese Dokumente im Tab „Beispiele hochladen" hoch.

**Was das System damit macht:**

1. **Text lesen** — Bei .docx-Dateien wird der Text direkt ausgelesen.
   Bei PDFs läuft OCR (wie bei den Büchern).

2. **Steckbrief** — Auch das Beispieldokument bekommt einen Steckbrief.
   Das System versteht also: *„Dieses Beispiel handelt von Lyrischen Texten."*

3. **Gespeichert** — Text + Steckbrief werden in einer eigenen Tabelle (`examples`) abgelegt.

**Was passiert beim nächsten Abfragen?**

Wenn du ein Thema abfragst (z.B. *„Lyrische Texte"*), passiert jetzt zusätzlich:
- Das System schaut: Gibt es ein Beispieldokument, das gut zu diesem Thema passt?
- Wenn ja (Ähnlichkeit ≥ 50 %), wird dieses Beispiel an Mistral mitgeschickt mit der Anweisung:
  *„Orientiere dich am Aufbau und Stil dieses Beispiels."*
- Mistral schreibt dann die Zusammenfassung im gewünschten Format — mit Content 1, Content 2 usw.

---

**❓ Ändert sich dabei der System-Prompt?**

**Nein.** Der System-Prompt (die Grundanweisung an Mistral) bleibt immer gleich.
Das Beispieldokument wird *extra hinzugefügt* — als Anhang im selben Paket wie die Schulbuch-Auszüge.
Mistral sieht also drei Dinge auf einmal:
- 📌 Das Thema
- 📖 Die 10 relevantesten Buchseiten
- 📄 Das passende Beispieldokument als Stilvorlage (wenn vorhanden)

Der System-Prompt ist editierbar unter ⚙️ im Tab „Thema abfragen" — Änderungen dort gelten sofort
und werden dauerhaft gespeichert.
        """)

    st.markdown("---")

    with st.expander("🎛️ Wie kann ich die Ergebnisse beeinflussen?"):
        st.markdown("""
Das System liefert zwei Outputs. Beide kannst du beeinflussen.

---

#### 📚 Output 1 — Textauszüge aus den Büchern

Das sind die Rohseiten aus den Schulbüchern, die das System als passend eingestuft hat.

**Was du tun kannst, um bessere Treffer zu bekommen:**

- **Mehr Kontext zum Thema geben** — Das System sucht anhand des Themennamens. Je mehr inhaltliche Beschreibung du zum Thema lieferst (z.B. aus dem Lehrplan: welche Aspekte sollen abgedeckt sein?), desto gezielter die Treffer. Das geht heute über die Themenformulierung im Dropdown — in Zukunft könnte man direkt Lehrplan-Text anhängen.
- **Mehr oder weniger Treffer** — Das Zahlenfeld rechts neben dem Dropdown (Standard: 10). Mehr Treffer = mehr Buchseiten im Ergebnis.
- **Andere Bücher einbeziehen** — Neue Bücher in Tab „Bücher hochladen" indexieren. Mehr Bücher = mehr mögliche Fundstellen.

---

#### 📝 Output 2 — Zusammenfassung

Das ist der Text, den die KI aus den gefundenen Buchseiten schreibt.

**Was du tun kannst, um eine bessere Zusammenfassung zu bekommen:**

- **Mehr Beispiele hochladen** — Je mehr fertige Texte du im Tab „Beispiele hochladen" hinterlegst, desto klarer versteht die KI, wie das Ergebnis aussehen soll. Format, Struktur, Tonalität — alles wird aus deinen Beispielen gelernt.
- **Den Auftrag präzisieren** — Über den System-Prompt (⚙️ im Tab „Thema abfragen") kannst du beschreiben, was die Zusammenfassung leisten soll. Zum Beispiel: *„Hilf Lehrern dabei, eine 45-Minuten-Stunde vorzubereiten"* oder *„Erkläre das Thema so, dass ein Berufsanfänger sofort loslegen kann."* Je klarer der Auftrag, desto besser das Ergebnis.
- **Neu generieren** — Der 🔄-Button erzeugt eine neue Version aus denselben Quellen. Gut zum Vergleichen nach einer Prompt-Änderung oder einem neuen Beispieldokument.
- **Modelle vergleichen** — In Zukunft: verschiedene KI-Modelle für die Zusammenfassung ausprobieren und die Ergebnisse nebeneinanderlegen. Heute läuft alles auf Mistral Large.
        """)

    st.markdown("---")

    with st.expander("💰 Was kostet was?"):
        st.markdown("""
Es gibt drei Kostenarten:

Alle Kosten entstehen bei **Mistral** (mistral.ai) — OCR, Steckbrief und Zusammenfassung laufen alle über Mistral-Modelle.

| | Wofür | Mistral-Modell | Wann | Kosten | Klett Deutsch (109 S.) |
|---|---|---|---|---|---|
| **1. OCR** | Buch lesen (Bild → Text) | `mistral-ocr-latest` | einmalig beim Hochladen | ~€0,002 pro Seite | €0,22 |
| **2. Steckbrief** | Text in Zahlen umwandeln | `mistral-embed` | einmalig beim Hochladen | ~€0,01 pro Buch | €0,004 |
| **3. Zusammenfassung** | KI schreibt den fertigen Text | `mistral-large-latest` | einmalig pro Thema | ~€0,01–0,02 | €0,01–0,02 |
| **Gesamt** | | | | | **~€0,23–0,24** |

Nach der ersten Runde kostet jede weitere Abfrage **€0** — alles ist gecacht.
        """)

    st.markdown("---")

    st.subheader('🎯 Wann ist das System "fertig"?')
    st.markdown("""
Das vertraglich vereinbarte Ziel:

> Für ein Lehrplan-Thema werden die **Top-10 Ergebnisse** angeschaut.
> Mindestens **8 von 10** müssen wirklich relevant sein.
> Das muss für **mindestens 5 verschiedene Themen** stabil funktionieren.

Sobald das erreicht ist, gilt das System als abgenommen. ✅
    """)

    st.markdown("---")

    with st.expander("📊 Was ist gerade in der Datenbank?"):
        st.markdown("""
Das ist der aktuelle Stand — was das System schon kennt und durchsuchen kann.

| Tabelle | Was drin ist |
|---|---|
| **Bücher** (`documents`) | 424 Seiten — Klett *Deutsch kompetent EF* (109 S.) + Paul D *Oberstufe Gesamtband* (315 S.), alle Fach: Deutsch |
| **Themen** (`topics`) | 93 Themen aus Excel (Deutsch + Mathe, EF/GK/LK) + extrahierte Kernlehrplan-Themen |
| **Zusammenfassungs-Cache** (`summary_cache`) | Wächst mit jeder Abfrage — Wiederholungen kosten €0 |
| **Einstellungen** (`settings`) | 2 System-Prompts (Extraktion + Zusammenfassung), editierbar in der App |
| **Beispiele** (`examples`) | Hochgeladene Beispieldokumente als Stilvorlagen |

**Erste Testergebnisse:**

| Thema | Bester Treffer | Ähnlichkeit |
|---|---|---|
| Sprachvarietäten | Seite 100 — Lexikon Sprache | 88% |
| Lyrische Texte | Seite 94 — Gattungslexikon Lyrik | 88% |
| Kommunikationsmodelle | Seite 6 — Kommunikationsmodelle | 88% |

Alle drei Testthemen landen sofort auf Platz 1. Der formale Abnahmetest (Top-10, ≥ 8/10) steht noch aus.
        """)

    with st.expander("🔧 Technischer Stack"):
        st.markdown("""
| Komponente | Technologie | Warum |
|---|---|---|
| **PDF lesen (OCR)** | Mistral OCR `mistral-ocr-latest` | Erkennt auch eingescannte Seiten zuverlässig |
| **Steckbriefe** | Mistral Embed `mistral-embed` (1024 Zahlen) | EU-gehostet, DSGVO-konform, günstig — alles bei einem Anbieter |
| **Datenbank** | Supabase + pgvector | Volle Datenkontrolle, skaliert gut |
| **KI-Zusammenfassung** | Mistral Large `mistral-large-latest` | Stark, EU-gehostet |
| **Oberfläche** | Streamlit | Schnell zu bauen, reicht für internes Tool |
| **Hosting** | Streamlit Cloud | Auto-Deploy bei jedem GitHub-Push |
| **Code** | GitHub `jansawatzki/studentpro-pdf-engine` | Public repo |

**Daten-Fluss in einem Satz:**
PDF → OCR (Text) → Embed (Zahlen) → Supabase (speichern) → bei Abfrage: Thema embedden → nächste Nachbarn suchen → Mistral Large zusammenfassen → Cache
        """)

